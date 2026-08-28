"""Testy ekstraktora plików DOCX."""

from __future__ import annotations

import io
import zipfile
from pathlib import Path

import docx
import pytest

from gnb.core.stale import PoziomPewnosciStruktury, RodzajBloku, TypZrodla
from gnb.core.wyjatki import BladTrwaly
from gnb.extractors.plik_docx import EkstraktorDocx

KATALOG_DANYCH = Path(__file__).resolve().parents[1] / "dane"


def _docx_z_przykladowa_trescia() -> bytes:
    """Buduje plik DOCX z nagłówkami, akapitem, listami i tabelą w kolejności."""
    dokument = docx.Document()
    dokument.add_heading("Tytuł dokumentu", level=1)
    dokument.add_paragraph("Wprowadzający akapit.")
    dokument.add_heading("Pierwszy podrozdział", level=2)
    dokument.add_paragraph("Pierwszy element", style="List Bullet")
    dokument.add_paragraph("Drugi element", style="List Bullet")
    dokument.add_paragraph("Trzeci element", style="List Bullet")
    dokument.add_heading("Drugi podrozdział", level=2)
    tabela = dokument.add_table(rows=2, cols=2)
    tabela.cell(0, 0).text = "Kolumna A"
    tabela.cell(0, 1).text = "Kolumna B"
    tabela.cell(1, 0).text = "1"
    tabela.cell(1, 1).text = "2"

    bufor = io.BytesIO()
    dokument.save(bufor)
    return bufor.getvalue()


def test_naglowki_listy_i_tabela_sa_rozpoznane() -> None:
    dokument = EkstraktorDocx().wyekstrahuj("plik_dokument-1", _docx_z_przykladowa_trescia())

    rodzaje = [blok.rodzaj for blok in dokument.bloki]
    assert rodzaje.count(RodzajBloku.NAGLOWEK) == 3
    assert RodzajBloku.LISTA in rodzaje
    assert RodzajBloku.TABELA in rodzaje
    assert dokument.poziom_pewnosci_struktury is PoziomPewnosciStruktury.WYSOKI


def test_elementy_listy_sa_sklejone_w_jeden_blok() -> None:
    dokument = EkstraktorDocx().wyekstrahuj("plik_dokument-2", _docx_z_przykladowa_trescia())

    bloki_listy = [blok for blok in dokument.bloki if blok.rodzaj is RodzajBloku.LISTA]
    assert len(bloki_listy) == 1
    assert bloki_listy[0].tresc.split("\n") == [
        "Pierwszy element",
        "Drugi element",
        "Trzeci element",
    ]


def test_tabela_zachowuje_wiersze_i_komorki() -> None:
    dokument = EkstraktorDocx().wyekstrahuj("plik_dokument-3", _docx_z_przykladowa_trescia())

    blok_tabeli = next(blok for blok in dokument.bloki if blok.rodzaj is RodzajBloku.TABELA)
    assert blok_tabeli.tresc.split("\n") == ["Kolumna A\tKolumna B", "1\t2"]


def test_tytul_bez_metadanych_pochodzi_z_pierwszego_naglowka() -> None:
    dokument = EkstraktorDocx().wyekstrahuj("plik_dokument-4", _docx_z_przykladowa_trescia())
    assert dokument.tytul == "Tytuł dokumentu"


def test_tytul_z_metadanych_ma_pierwszenstwo() -> None:
    surowy = docx.Document()
    surowy.add_heading("Nagłówek w treści", level=1)
    surowy.core_properties.title = "Tytuł z właściwości pliku"
    bufor = io.BytesIO()
    surowy.save(bufor)

    dokument = EkstraktorDocx().wyekstrahuj("plik_dokument-5", bufor.getvalue())
    assert dokument.tytul == "Tytuł z właściwości pliku"


def test_pusty_dokument_nie_ma_blokow_ani_tytulu() -> None:
    pusty = docx.Document()
    bufor = io.BytesIO()
    pusty.save(bufor)

    dokument = EkstraktorDocx().wyekstrahuj("plik_dokument-6", bufor.getvalue())
    assert dokument.bloki == []
    assert dokument.tytul is None
    assert dokument.tekst == ""


def test_obsluguje_wylacznie_format_docx() -> None:
    ekstraktor = EkstraktorDocx()
    assert ekstraktor.obsluguje(TypZrodla.PLIK_DOKUMENT, "docx") is True
    assert ekstraktor.obsluguje(TypZrodla.PLIK_DOKUMENT, "pdf") is False


def test_plik_testowy_ma_naglowki_akapity_i_liste() -> None:
    dane = (KATALOG_DANYCH / "dokument.docx").read_bytes()
    dokument = EkstraktorDocx().wyekstrahuj("plik_dokument-7", dane)

    rodzaje = [blok.rodzaj for blok in dokument.bloki]
    assert RodzajBloku.NAGLOWEK in rodzaje
    assert RodzajBloku.AKAPIT in rodzaje
    assert RodzajBloku.LISTA in rodzaje
    assert dokument.tytul == "Jak przygotować bazę wiedzy dla asystenta AI"


def _archiwum_zip_bez_dokumentu() -> bytes:
    """Buduje poprawne archiwum zip, które nie jest dokumentem DOCX."""
    bufor = io.BytesIO()
    with zipfile.ZipFile(bufor, "w") as archiwum:
        archiwum.writestr("cokolwiek.txt", "to nie jest dokument")
    return bufor.getvalue()


def test_uszkodzony_plik_konczy_sie_bledem_trwalym() -> None:
    """Plik, który nie jest archiwum zip, nie może wywrócić całego przebiegu.

    Biblioteka zgłasza tu wyjątek spoza taksonomii projektu, więc potok go nie
    łapał i poprawne źródła z tej samej partii nie były przetwarzane.
    """
    with pytest.raises(BladTrwaly, match="uszkodzony"):
        EkstraktorDocx().wyekstrahuj("plik_dokument-20", b"to nie jest plik DOCX")


def test_archiwum_bez_zawartosci_dokumentu_konczy_sie_bledem_trwalym() -> None:
    """Poprawne archiwum zip bez pliku Content_Types też jest błędem trwałym."""
    with pytest.raises(BladTrwaly, match="uszkodzony"):
        EkstraktorDocx().wyekstrahuj("plik_dokument-21", _archiwum_zip_bez_dokumentu())
